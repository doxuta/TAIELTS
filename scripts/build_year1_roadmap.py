# -*- coding: utf-8 -*-
"""Generate [01] Lộ Trình Chi Tiết 12 Tháng.docx — Năm 1 nền tảng tiếng Anh thuần.
Năm 2 để mở — chốt định hướng (IELTS / pháp lý / học thuật / giao tiếp) sau Q3/Q4."""
from pathlib import Path

from docx import Document
from docx.shared import Pt

BASE = Path(__file__).resolve().parent.parent
OUT = BASE / "01 - Lộ Trình & Kế Hoạch" / "📅 [01] Lộ Trình Chi Tiết 12 Tháng.docx"


def hr(doc: Document):
    p = doc.add_paragraph("━" * 40)
    p.runs[0].font.size = Pt(9)


def h(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def p(doc: Document, text: str, bold: bool = False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    return para


def bullet(doc: Document, text: str):
    doc.add_paragraph(text, style="List Bullet")


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    h(doc, "[01] LỘ TRÌNH CHI TIẾT — NĂM 1: NỀN TẢNG TIẾNG ANH (12 THÁNG)", 0)
    p(doc, "Học viên: Huyen Trang | Phụ thuộc: [MASTER] Hồ Sơ Học Viên (trình độ đầu vào, lịch học).")
    p(doc, "Cập nhật khi: có kết quả bài test đầu vào; review bắt buộc cuối mỗi quý để điều chỉnh tốc độ.")
    hr(doc)

    p(
        doc,
        "Nguyên tắc Năm 1: lấy GỐC TIẾNG ANH làm trọng tâm — không ép IELTS, không ép band theo tháng. "
        "Ưu tiên năng lực nền: ngữ pháp dùng đúng, phát âm chuẩn các âm khó, nghe hiểu đoạn ngắn–vừa, "
        "nói liên tục 1–2 phút, đọc hiểu đoạn 200–350 từ, viết câu–đoạn có kiểm soát. "
        "Năm 2 sẽ chốt định hướng sau (IELTS / tiếng Anh pháp lý / học thuật / giao tiếp) — phụ thuộc tiến độ + nhu cầu thực tế của HV cuối Năm 1.",
    )
    hr(doc)

    h(doc, "Khung 3 buổi / tuần (mỗi buổi 120 phút)", 1)
    p(doc, "Lặp lại hàng tuần; GV có thể hoán đổi thứ tự A/B/C theo lịch nhưng mỗi tuần phải đủ 1A + 1B + 1C.")
    bullet(
        doc,
        "Buổi A — Ngữ pháp + ứng dụng + từ vựng: ~60' chủ đề ngữ pháp (giải thích ngắn + 20–30 câu drill có đáp án) "
        "+ ~20' speaking ứng dụng đúng cấu trúc vừa học (đặt câu, mô tả ngắn) + ~40' ôn/mở rộng từ vựng chủ đề tuần "
        "(flashcard, đặt câu, collocation). Tránh để HV ngồi nghe lý thuyết quá 60 phút liên tục.",
    )
    bullet(
        doc,
        "Buổi B — Nghe + Nói + Phát âm: ~45' listening (nghe có transcript, gap-fill, true/false đơn giản) "
        "+ ~50' speaking (câu hỏi gợi ý, mô tả tranh/chủ đề quen, ghi âm tự nghe lại) + ~15' chuyên đề phát âm "
        "(âm khó, word/sentence stress) + ~10' nhận xét/ghi lỗi vào Sổ Tay.",
    )
    bullet(
        doc,
        "Buổi C — Đọc + Viết nền: ~50' reading (đoạn 200–350 từ, câu hỏi detail/main idea/vocab-in-context) "
        "+ ~50' writing (câu ghép, đoạn 80–150 từ theo khung — topic sentence + supporting) + ~20' từ vựng/collocation.",
    )
    hr(doc)

    h(doc, "Bài kiểm tra cuối tháng (Năm 1 — thuần năng lực nền)", 1)
    bullet(doc, "Tháng 1–11: Kiểm tra nền — trắc nghiệm/tự luận ngắn ngữ pháp + 1 đoạn nghe + 1 bài đọc + viết 80–150 từ + nói 2 phút. KHÔNG dùng format IELTS.")
    bullet(doc, "Tháng 12: Kiểm tra tổng hợp năm + rubric năng lực nền theo 6 kỹ năng (ngữ pháp / từ vựng / nghe / nói / đọc / viết). Dùng kết quả này để chốt định hướng Năm 2.")
    hr(doc)

    # Pronunciation focus
    h(doc, "Chuyên đề phát âm xuyên suốt Năm 1 (15 phút/buổi B)", 1)
    bullet(doc, "Tháng 1: âm cuối /s/-/z/-/ɪz/ (số nhiều, sở hữu, chia thì hiện tại đơn ngôi thứ 3).")
    bullet(doc, "Tháng 2: âm cuối -ed (/t/-/d/-/ɪd/) sau động từ quá khứ đơn.")
    bullet(doc, "Tháng 3: word stress 2–3 âm tiết (record n/v, contract n/v, present n/v…).")
    bullet(doc, "Tháng 4–6: cặp âm dễ lẫn (/ɪ/ vs /iː/, /æ/ vs /e/, /θ/ vs /s/, /ʃ/ vs /tʃ/).")
    bullet(doc, "Tháng 7–9: sentence stress + content vs function words; linking sounds đơn giản.")
    bullet(doc, "Tháng 10–12: intonation câu khẳng định/hỏi/liệt kê; rhythm trong đoạn nói 2 phút.")
    p(doc, "Mỗi buổi B mở đầu/khép lại bằng 5–10 phút drill phát âm; ghi lỗi vào Sổ Tay Lỗi Sai mục Phát Âm.")

    # Quarters
    months_data = [
        # Q1
        (
            1,
            "THÁNG 1 — XÂY NỀN CÂU & THÌ CƠ BẢN",
            [
                "Tuần 1: Câu khẳng định/phủ định/câu hỏi Yes-No & Wh-; bài tập: 30 câu chia loại câu + 10 câu sắp xếp trật tự từ.",
                "Tuần 2: Hiện tại đơn (dùng, dấu hiệu) + tần suất (always/usually…); bài tập: 25 câu chia thì + 1 đoạn chọn dạng đúng.",
                "Tuần 3: Hiện tại tiếp diễn vs đơn; bài tập: 30 câu + viết 8 câu mô tả hành động đang diễn ra.",
                "Tuần 4: Quá khứ đơn (động từ bất quy tắc top 20); bài tập: 25 câu + kể 1 sự kiện quá khứ (80–100 từ).",
            ],
            "Từ vựng: Gia đình, công việc, học tập, sinh hoạt hằng ngày (~50 từ/tuần, ưu tiên collocations gốc).",
            "Phát âm tháng 1: /s/-/z/-/ɪz/ (works/plays/watches) — drill 5 phút mỗi buổi B.",
            "Kiểm tra tháng 1: Ngữ pháp cốt lõi + nghe 1 đoạn 2–3 phút + đọc 1 đoạn 200 từ + viết đoạn 80–100 từ + nói 90 giây.",
        ),
        (
            2,
            "THÁNG 2 — THÌ NÂNG CAO & LIÊN KẾT CÂU",
            [
                "Tuần 1: Quá khứ tiếp diễn; bài tập: 25 câu + miêu tả 1 khoảnh khắc trong quá khứ.",
                "Tuần 2: Hiện tại hoàn thành (kết nối với quá khứ); bài tập: 30 câu + timeline cá nhân (viết).",
                "Tuần 3: Tương lai (will / be going to / present continuous for arrangements); bài tập: 28 câu + lịch cá nhân tuần tới (viết+nói).",
                "Tuần 4: Ôn tổ hợp thì — bài tập mixed 40 câu + speaking 3 chủ đề quen (2 phút/chủ đề).",
            ],
            "Từ vựng: Sức khỏe, thói quen, thời tiết, du lịch cơ bản.",
            "Phát âm tháng 2: -ed endings (/t/-/d/-/ɪd/) — walked/played/wanted.",
            "Kiểm tra tháng 2: Mixed grammar + listening gap-fill + reading detail/main idea + viết 100–120 từ + nói 90 giây.",
        ),
        (
            3,
            "THÁNG 3 — CÂU PHỨC CƠ BẢN & ĐỌC HIỂU DÀI HƠN",
            [
                "Tuần 1: Liên từ (and/but/because/so); ghép câu; bài tập: 20 cặp câu ghép + 1 đoạn chèn liên từ.",
                "Tuần 2: Mệnh đề trạng ngữ thời gian (when/before/after); bài tập: 25 câu + viết lịch trình 1 ngày.",
                "Tuần 3: Mệnh đề danh từ (that/if/whether) giới thiệu; bài tập: 18 câu + nói 90 giây về 1 quyết định.",
                "Tuần 4: Ôn Q1 — checklist thì + liên kết; project nhỏ: 'Giới thiệu bản thân 2 phút' + bản viết 120 từ.",
            ],
            "Từ vựng: Công nghệ đời thường, học tập, môi trường xung quanh.",
            "Phát âm tháng 3: word stress 2–3 âm tiết (REcord vs reCORD; CONtract vs conTRACT).",
            "Kiểm tra tháng 3 (cuối Q1): 4 kỹ năng nền — rubric chính xác ngữ pháp / mạch lạc / hiểu đoạn / phát âm rõ.",
        ),
        # Q2
        (
            4,
            "THÁNG 4 — BỊ ĐỘNG & CÂU ĐIỀU KIỆN LOẠI 0–1",
            [
                "Tuần 1: Passive cơ bản (hiện tại & quá khứ đơn); bài tập: 22 câu đổi thể chủ động/bị động.",
                "Tuần 2: Passive với modal; bài tập: 18 câu + viết 3 câu mô tả quy trình (luật/học tập).",
                "Tuần 3: Conditional 0 & 1; bài tập: 24 câu + viết 5 'If…' thực tế.",
                "Tuần 4: Reading dài 250–320 từ + câu hỏi detail; speaking: mô tả 1 quy trình 90 giây.",
            ],
            "Từ vựng: Chủ đề xã hội đơn giản (traffic, public services, study habits).",
            "Phát âm tháng 4: cặp /ɪ/ vs /iː/ (ship/sheep; live/leave; bit/beat).",
            "Kiểm tra tháng 4: Grammar focus + 1 reading + viết đoạn 100–130 từ có dùng passive/conditional.",
        ),
        (
            5,
            "THÁNG 5 — MỆNH ĐỀ QUAN HỆ & ĐẠI TỪ QUAN HỆ",
            [
                "Tuần 1: who/which/that trong mệnh đề xác định; bài tập: 20 câu nối câu + 1 đoạn có lỗi sửa.",
                "Tuần 2: whose/where/when; bài tập: 18 câu + mô tả nơi chốn quen thuộc.",
                "Tuần 3: Mệnh đề không xác định (which/who, dấu phẩy); bài tập: 16 câu + viết 120 từ giới thiệu 1 người.",
                "Tuần 4: Speaking: 3 phần mô tả ngắn — warm-up, 1 phút chuẩn bị, 2 phút nói (chủ đề đời thường, không IELTS Part 2).",
            ],
            "Từ vựng: Con người, địa điểm, sự kiện; học cụm động từ + giới từ.",
            "Phát âm tháng 5: cặp /æ/ vs /e/ (bad/bed; man/men; sat/set).",
            "Kiểm tra tháng 5: Relative clauses + listening (note-taking 5 bullet) + viết 120–140 từ.",
        ),
        (
            6,
            "THÁNG 6 — CÂU ƯỚC & MODAL CƠ BẢN + VIẾT ĐOẠN LIỀN MẠCH",
            [
                "Tuần 1: wish/if only (hiện tại đơn giản); bài tập: 15 câu + nói 3 điều muốn thay đổi (90 giây).",
                "Tuần 2: can/could/may/might (khả năng & xin phép); bài tập: 20 tình huống chọn modal.",
                "Tuần 3: must/have to/should (nghĩa vụ & lời khuyên); bài tập: 18 câu + thư email ngắn (100 từ).",
                "Tuần 4: Ôn Q2 — viết đoạn 130–150 từ có topic sentence + 2 supporting sentences + ví dụ.",
            ],
            "Từ vựng: Cảm xúc, học tập, mục tiêu cá nhân; bắt đầu collocation học thuật nhẹ (10 cụm/tuần).",
            "Phát âm tháng 6: cặp /θ/ vs /s/ (think/sink; thick/sick) — lưỡi giữa răng.",
            "Kiểm tra tháng 6 (cuối Q2): Kiểm tra nền + speaking 2 phút không script + rubric mạch ý.",
        ),
        # Q3
        (
            7,
            "THÁNG 7 — GERUND / INFINITIVE & REPORTED SPEECH GIỌT ĐẦU",
            [
                "Tuần 1: Gerund vs infinitive (verb patterns phổ biến); bài tập: 30 câu chọn dạng.",
                "Tuần 2: Verb + object + infinitive; bài tập: 20 câu + viết 5 câu về lời khuyên.",
                "Tuần 3: Reported speech (thì lùi cơ bản); bài tập: 16 câu chuyển lời nói trực tiếp/gián tiếp.",
                "Tuần 4: Listening: hội thoại 3–4 phút (2 lần nghe) + trả lời 8 câu hỏi.",
            ],
            "Từ vựng: Giao tiếp, tin tức đời sống; nhấn word formation đơn giản.",
            "Phát âm tháng 7: sentence stress — nhấn content word, lướt function word.",
            "Kiểm tra tháng 7: Reported speech + listening Q&A + reading 300 từ.",
        ),
        (
            8,
            "THÁNG 8 — SO SÁNH & TRẠNG TỪ + ĐỌC PHÂN TÍCH NHẸ",
            [
                "Tuần 1: So sánh hơn/kém nhất; bài tập: 25 câu + mô tả 2 đồ vật (viết).",
                "Tuần 2: So sánh kép / as…as; bài tập: 18 câu + speaking so sánh 2 thói quen học.",
                "Tuần 3: Trạng từ tần suất & mức độ; bài tập: 20 câu + paragraph 130 từ.",
                "Tuần 4: Reading: skim main idea + scan detail trên đoạn 300–400 từ.",
            ],
            "Từ vựng: Mô tả thay đổi/xu hướng đơn giản (rise/fall/remain) — phục vụ viết mô tả, không chuyên Task 1.",
            "Phát âm tháng 8: content vs function word stress + linking phụ âm-nguyên âm.",
            "Kiểm tra tháng 8: So sánh + adverbs + viết mô tả ngắn 120–150 từ + đọc 300 từ.",
        ),
        (
            9,
            "THÁNG 9 — LIÊN KẾT ĐOẠN & NÓI DÀI HƠN",
            [
                "Tuần 1: Linkers (however/therefore/because of/in addition); bài tập: chèn linker vào đoạn thiếu.",
                "Tuần 2: Problem–solution khung 4 câu (nền); viết 150 từ chủ đề quen.",
                "Tuần 3: Speaking: 3 phút có gợi ý bullet; ghi âm tự nghe lại; ghi chú phát âm.",
                "Tuần 4: Ôn Q3 — portfolio: 3 bài viết ngắn đã sửa + checklist lỗi hay gặp.",
            ],
            "Từ vựng: Chủ đề học thuật nhẹ (education, health basics, environment).",
            "Phát âm tháng 9: intonation câu hỏi yes/no vs câu hỏi wh-.",
            "Kiểm tra tháng 9 (cuối Q3): 4 kỹ năng nền — độ khó vừa; nhấn feedback cá nhân hơn điểm số. GV bàn với HV về định hướng Năm 2.",
        ),
        # Q4 — Năm 1 thuần nền, không IELTS-lite
        (
            10,
            "THÁNG 10 — ĐỌC PHỨC TẠP HƠN & NGHE DÀI HƠN",
            [
                "Tuần 1: Reading 350–450 từ; câu hỏi paragraph purpose + vocabulary in context.",
                "Tuần 2: Listening: đoạn 4–5 phút, 1 lần nghe + note outline 6 ý.",
                "Tuần 3: Viết opinion 180–200 từ (4 đoạn — intro/2 ý chính/kết).",
                "Tuần 4: Speaking: phản biện nhẹ 2 vế (agree/disagree đơn giản) 4 phút có gợi ý.",
            ],
            "Từ vựng: 40 từ chủ đề ý kiến/lập luận (in my view, on the other hand, evidence shows…).",
            "Phát âm tháng 10: intonation liệt kê + ngắt câu dài tự nhiên.",
            "Kiểm tra tháng 10: Bài kiểm tra nền 4 kỹ năng — đọc 400 từ + nghe 4 phút + viết 180 từ + nói 3 phút.",
        ),
        (
            11,
            "THÁNG 11 — TỔNG HỢP KỸ NĂNG NỀN",
            [
                "Tuần 1: Listening: phỏng vấn ngắn 4–5 phút + 10 câu hỏi (note-taking, T/F, MCQ — không cố định IELTS section).",
                "Tuần 2: Reading: 2 đoạn liền chủ đề (so sánh ý kiến) + 10 câu hỏi.",
                "Tuần 3: Viết: 1 đoạn mô tả 1 dữ liệu đơn giản (biểu đồ cột hoặc bảng — phục vụ kỹ năng viết mô tả, không bắt buộc theo Task 1 IELTS).",
                "Tuần 4: Speaking: bài nói 2 phút có gợi ý bullet về chủ đề học/làm việc; theo sau là Q&A 2 phút.",
            ],
            "Từ vựng: Collocation học thuật chung; có thể thêm 10 cụm Luật song ngữ nếu HV yêu cầu (tùy chọn).",
            "Phát âm tháng 11: rhythm + tự nghe lại 1 bài 2 phút và đánh dấu lỗi.",
            "Kiểm tra tháng 11: Tổng hợp 4 kỹ năng + nói 3 phút — rubric nền, không quy đổi band.",
        ),
        (
            12,
            "THÁNG 12 — TỔNG HỢP NĂM 1 & CHỐT ĐỊNH HƯỚNG NĂM 2",
            [
                "Tuần 1: Ôn ngữ pháp theo lỗi cá nhân (từ Sổ Tay Lỗi Sai) — chữa lại 30 câu sai phổ biến nhất.",
                "Tuần 2: Tổng ôn đọc–viết: 1 bài tổng hợp 200 từ có outline GV duyệt + 1 đoạn đọc 400 từ.",
                "Tuần 3: Kiểm tra cuối năm (4 kỹ năng + phát âm, chia 2 buổi nếu cần).",
                "Tuần 4: Họp review với HV: rà soát rubric năng lực + cùng chốt định hướng Năm 2 (IELTS / tiếng Anh pháp lý / học thuật chung / giao tiếp) → cập nhật MASTER.",
            ],
            "Từ vựng: Ôn 200 từ hay sai + 50 collocation đã học.",
            "Phát âm tháng 12: bài tổng hợp 3 phút đọc thành tiếng + ghi âm — đối chiếu lỗi từ T1-T11.",
            "Kiểm tra tháng 12: Bài tổng kết năm + bảng tự đánh giá năng lực nền (rubric 6 kỹ năng × thang 1–5). Kết quả này quyết định nội dung Năm 2.",
        ),
    ]

    for mnum, title, weeks, vocab, pron, assessment in months_data:
        hr(doc)
        h(doc, f"THÁNG {mnum} — {title.split(' — ')[1]}", 1)
        p(doc, title)
        h(doc, "Ngữ pháp & bài tập (theo tuần)", 2)
        for w in weeks:
            bullet(doc, w)
        p(doc, vocab, bold=True)
        p(doc, pron, bold=True)
        h(doc, "Gợi ý bài tập về nhà hằng ngày (HV)", 2)
        bullet(doc, "15–20 phút Anki / danh sách từ (theo chủ đề tuần).")
        bullet(doc, "20–30 phút bài tập ngữ pháp (số lượng GV giao cụ thể trong giáo án).")
        bullet(doc, "15–25 phút nghe 1 đoạn + chép chính tả 1–2 câu hoặc trả lời 3 câu hỏi.")
        bullet(doc, "5–10 phút drill phát âm chủ đề tháng (đọc theo audio mẫu + ghi âm tự nghe).")
        h(doc, "Kiểm tra cuối tháng", 2)
        p(doc, assessment)

    hr(doc)
    h(doc, "NĂM 2 — ĐỂ MỞ, CHỐT ĐỊNH HƯỚNG CUỐI NĂM 1", 1)
    p(
        doc,
        "Năm 2 KHÔNG bắt buộc là IELTS. Sau khi hoàn thành nền 12 tháng và rà soát kết quả tháng 12, "
        "GV + HV cùng chốt 1 trong các hướng — và bổ sung chi tiết lộ trình Năm 2 vào file này:",
    )
    bullet(doc, "Hướng A — IELTS Academic: đăng ký thi 6.0–6.5 phục vụ hồ sơ Thạc sĩ Luật. Lộ trình 9–12 tháng làm quen format + mock.")
    bullet(doc, "Hướng B — Tiếng Anh Pháp Lý (Legal English): từ vựng/cấu trúc chuyên ngành, đọc văn bản luật, viết memo. Phù hợp nếu HV cần dùng tiếng Anh trong nghề trước khi thi quốc tế.")
    bullet(doc, "Hướng C — Tiếng Anh Học Thuật chung: viết essay, bài luận, đọc paper. Phù hợp nếu HV chưa rõ thời điểm thi nhưng muốn dùng tiếng Anh đọc tài liệu.")
    bullet(doc, "Hướng D — Giao tiếp/Du lịch nâng cao: nếu nhu cầu chính là giao tiếp công việc + đời sống. Có thể ghép với hướng A/B khi cần.")
    bullet(doc, "Quyết định dựa trên: kết quả Q4 + nhu cầu thực tế của HV + thời điểm hồ sơ Thạc sĩ. KHÔNG ép chọn IELTS chỉ vì mục tiêu ban đầu — luôn rà lại bối cảnh thực.")
    p(doc, "Luồng dữ liệu giữ nguyên: Test đầu vào → MASTER → Lộ trình → Giáo án → Sheet buổi/tuần/tháng → Báo cáo.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    main()
