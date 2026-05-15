# -*- coding: utf-8 -*-
"""Regenerate MASTER Hồ Sơ với mục tiêu Năm 1 nền + review quý."""
from pathlib import Path

from docx import Document
from docx.shared import Pt

BASE = Path(__file__).resolve().parent.parent
OUT = BASE / "01 - Lộ Trình & Kế Hoạch" / "📋 [01] MASTER - Hồ Sơ Học Viên Huyen Trang.docx"


def hr(doc):
    doc.add_paragraph("━" * 40)


def main():
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading("[MASTER FILE] HỒ SƠ HỌC VIÊN — HUYEN TRANG", 0)
    doc.add_paragraph("File nguồn gốc — mọi file khác tham chiếu từ đây.")
    hr(doc)

    doc.add_heading("A. THÔNG TIN CÁ NHÂN", level=1)
    doc.add_paragraph("Họ tên:           Huyen Trang")
    doc.add_paragraph("Ngày sinh:        [___________]")
    doc.add_paragraph("Nghề nghiệp:      Sinh viên Luật")
    doc.add_paragraph("Số điện thoại:    [___________]")
    doc.add_paragraph("Email:            [___________]")
    doc.add_paragraph("Kênh liên lạc:   Slack / Zalo")

    doc.add_heading("B. MỤC TIÊU HỌC TẬP (NĂM 1 = NỀN TẢNG)", level=1)
    doc.add_paragraph(
        "Năm 1 không lấy band IELTS làm mục tiêu chính. Band (nếu có) chỉ ghi sau test đầu vào "
        "hoặc bài tham chiếu cuối năm — dùng để theo dõi, không ép mốc theo tháng."
    )
    doc.add_paragraph("Mục tiêu ngắn hạn (3 tháng — hết Quý 1 Năm 1):")
    doc.add_paragraph(
        "→ Ổn các thì cơ bản + câu hỏi/câu phủ định; viết đoạn 80–120 từ có kiểm soát.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "→ Nghe hiểu đoạn 2–3 phút có transcript; nói liên tục 90 giây–2 phút về chủ đề quen.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "→ ~600 từ/chủ đề nền + một số collocation (ghi trong Danh sách từ vựng).",
        style="List Bullet",
    )

    doc.add_paragraph("Mục tiêu trung hạn (6 tháng — hết Quý 2 Năm 1):")
    doc.add_paragraph(
        "→ Passive & conditional 0–1; mệnh đề quan hệ; đọc đoạn 250–350 từ với câu hỏi detail.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "→ Viết đoạn 130–150 từ có topic sentence; nói mô tả/chuỗi ý ~2 phút.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "→ Chưa ưu tiên full format IELTS; tập trung năng lực nền (ngữ pháp, hội thoại, đọc hiểu).",
        style="List Bullet",
    )

    doc.add_paragraph("Mục tiêu 12 tháng (hết Năm 1):")
    doc.add_paragraph(
        "→ Hoàn thành nội dung file [01] Lộ Trình Chi Tiết 12 Tháng (Năm 1 nền tiếng Anh thuần).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "→ Đạt checklist năng lực nền cuối năm: 6 kỹ năng (ngữ pháp / từ vựng / nghe / nói / đọc / viết) "
        "đều đạt rubric nền theo bài kiểm tra tháng 12.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "→ Định hướng Năm 2 ĐỂ MỞ — GV + HV cùng chốt cuối tháng 12 dựa trên kết quả thực + nhu cầu HV. "
        "Các hướng khả thi: (A) IELTS Academic phục vụ Thạc sĩ Luật; (B) Tiếng Anh Pháp Lý; "
        "(C) Tiếng Anh học thuật chung; (D) Giao tiếp/làm việc. KHÔNG ép IELTS chỉ vì mục tiêu ban đầu.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Cam kết điều chỉnh: Review cuối mỗi quý (tháng 3, 6, 9, 12) để chỉnh tốc độ và "
        "trọng tâm ngữ pháp theo lỗi thực tế (Sổ Tay Lỗi Sai + Sheet buổi học). "
        "Định hướng Năm 2 rà thử ở cuối Q3 (tháng 9), chốt chính thức cuối Q4 (tháng 12)."
    )

    doc.add_heading("C. TRÌNH ĐỘ ĐẦU VÀO (điền sau test)", level=1)
    doc.add_paragraph("Ngày test:          [___________]")
    doc.add_paragraph("Ngữ pháp:           ___/20")
    doc.add_paragraph("Đọc hiểu:           ___/20")
    doc.add_paragraph("Viết:               ___/20")
    doc.add_paragraph("Nói:                ___/20")
    doc.add_paragraph("Tổng:               ___/80")
    doc.add_paragraph("Band tương đương (tham chiếu, không ép mốc Năm 1):   ___")
    doc.add_paragraph("Điểm mạnh:          [___________]")
    doc.add_paragraph("Điểm yếu ưu tiên:   [___________]")

    doc.add_heading("D. LỊCH HỌC", level=1)
    doc.add_paragraph("Số buổi/tuần:     3 buổi")
    doc.add_paragraph("Thời lượng:       2 tiếng/buổi")
    doc.add_paragraph("Địa điểm:         Quán cà phê")
    doc.add_paragraph("Các ngày học:     Thứ ___, ___, ___")
    doc.add_paragraph("Giờ học:          ___:___ – ___:___")
    doc.add_paragraph("Ngày bắt đầu:     [___________]")
    doc.add_paragraph("Ngày thi (nếu chọn IELTS Năm 2): [chốt cuối Q4 Năm 1 — không đặt mặc định]")

    doc.add_heading("E. HỌC PHÍ", level=1)
    doc.add_paragraph("Hình thức:       Theo tháng (12 buổi)")
    doc.add_paragraph("Mức phí:         ___________đ / tháng")
    doc.add_paragraph("Thanh toán:      Đầu mỗi tháng")
    doc.add_paragraph("Tiền nước:       [Chia đôi / Học viên trả]")

    doc.add_heading("F. LINKS CÁC FILE LIÊN KẾT", level=1)
    doc.add_paragraph("→ Lộ trình chi tiết:        [dán link vào đây]")
    doc.add_paragraph("→ Sheet theo dõi buổi học:  [dán link vào đây]")
    doc.add_paragraph("→ Sheet tổng kết tuần:      [dán link vào đây]")
    doc.add_paragraph("→ Sheet đánh giá tháng:     [dán link vào đây]")
    doc.add_paragraph("→ Danh sách từ vựng:        [dán link vào đây]")
    doc.add_paragraph("→ Ghi chú ngữ pháp:         [dán link vào đây]")
    doc.add_paragraph("→ Bài test đầu vào:         [dán link vào đây]")
    doc.add_paragraph("→ Template giáo án:         [dán link vào đây]")

    doc.add_heading("G. GHI CHÚ GIÁO VIÊN", level=1)
    doc.add_paragraph("Tính cách học viên:     [___________]")
    doc.add_paragraph("Phong cách học:         [___________]")
    doc.add_paragraph("Điều cần lưu ý:         [___________]")
    doc.add_paragraph("Cập nhật lần cuối:      [___________]")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    main()
