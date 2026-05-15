# -*- coding: utf-8 -*-
"""Rebuild lesson/homework/assessment templates for Năm 1 nền (rubric nền, buổi A/B/C)."""
from pathlib import Path

from docx import Document
from docx.shared import Pt

BASE = Path(__file__).resolve().parent.parent


def new_doc():
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)
    return doc


def add_block(doc, title: str, lines: list[str]):
    doc.add_heading(title, level=1)
    for line in lines:
        doc.add_paragraph(line)


def build_giao_an():
    doc = new_doc()
    doc.add_heading("[TEMPLATE] GIÁO ÁN TỪNG BUỔI HỌC — NĂM 1 NỀN", 0)
    doc.add_paragraph(
        "Phụ thuộc: Lộ Trình Chi Tiết (tuần/tháng); Danh Sách Từ Vựng; Sheet Theo Dõi Buổi Học (sau buổi)."
    )
    doc.add_paragraph(
        'Hướng dẫn: Copy file → đổi tên "Giáo Án - Buổi [số] - [ngày]". Mỗi tuần cần đủ 1 buổi A + 1 B + 1 C (có thể đổi thứ tự theo lịch).'
    )
    add_block(
        doc,
        "THÔNG TIN BUỔI HỌC",
        [
            "Buổi số:        ___",
            "Ngày:           ___/___/______",
            "Giờ:            ___:___ – ___:___",
            "Địa điểm:       ___________",
            "Tuần số:        ___  |  Tháng:  ___",
            "Loại buổi (chọn một):",
            "  [ ] A — Ngữ pháp (60') + speaking ứng dụng (20') + từ vựng/ôn (40')",
            "  [ ] B — Listening (45') + speaking (50') + chuyên đề phát âm (15') + nhận xét (10')",
            "  [ ] C — Reading (50') + writing (50') + từ vựng/collocation (20')",
            "Chủ đề ngữ pháp tuần (từ lộ trình): ___________________________",
            "Số bài tập ngữ pháp giao (ước lượng): ___ câu / ___ trang _________",
        ],
    )
    add_block(
        doc,
        "PHÂN BỔ THỜI LƯỢNG TRONG BUỔI (điền phút — tổng ~120)",
        [
            "Ngữ pháp + chữa bài tập:     _____ phút  (Buổi A: ~60')",
            "Speaking ứng dụng/giao tiếp: _____ phút  (Buổi A: ~20' | Buổi B: ~50')",
            "Listening:                   _____ phút  (Buổi B: ~45')",
            "Chuyên đề phát âm:           _____ phút  (Buổi B: ~15')",
            "Reading:                     _____ phút  (Buổi C: ~50')",
            "Writing (câu/đoạn):          _____ phút  (Buổi C: ~50')",
            "Ôn từ vựng / collocation:    _____ phút  (Buổi A: ~40' | Buổi C: ~20')",
            "Nhận xét & giao bài về nhà:  _____ phút",
        ],
    )
    add_block(
        doc,
        "PHẦN 1 — WARM UP (gợi ý 10–15 phút)",
        [
            "Kiểm tra bài tập buổi trước:  [ ] Có  [ ] Không  [ ] Nộp muộn",
            "Chất lượng:  [ ] Tốt  [ ] TB  [ ] Cần cải thiện",
            "Ôn nhanh 5 từ (random từ Danh Sách): ghi từ + ✓/✗",
        ],
    )
    add_block(
        doc,
        "PHẦN 2 — NỘI DUNG CHÍNH (theo loại buổi A / B / C)",
        [
            "Buổi A (60'+20'+40'): Giải thích ngắn chủ đề ngữ pháp → 20-30 câu bài tập có đáp án → chữa lỗi điển hình (ghi Sổ Tay) → 20' speaking ứng dụng đúng cấu trúc vừa học (đặt câu/mô tả ngắn) → 40' ôn/mở rộng từ vựng chủ đề tuần.",
            "Buổi B (45'+50'+15'+10'): Nghe 1-2 đoạn có transcript → gap-fill/Q&A → Speaking có gợi ý bullet, ghi âm tự nghe → 15' chuyên đề phát âm tháng (xem lộ trình) → nhận xét và ghi lỗi vào Sổ Tay.",
            "Buổi C (50'+50'+20'): Reading 1 đoạn 200-400 từ theo tháng → câu hỏi detail/main idea/vocab-in-context → Writing câu ghép hoặc đoạn 80-150 từ có khung (topic sentence + supporting) → 20' từ vựng/collocation chủ đề tuần.",
            "Ghi chú nội dung đã làm: _____________________________________________",
        ],
    )
    add_block(
        doc,
        "PHẦN 3 — TỔNG KẾT BUỔI (5–10 phút)",
        [
            "Điểm mạnh hôm nay: ___________________________",
            "Cần khắc phục:     ___________________________",
            "Bài tập về nhà (ghi rõ số câu/trang): ___________________________",
        ],
    )
    add_block(
        doc,
        "PHẦN 4 — SAU BUỔI (GV điền)",
        [
            "Đánh giá thái độ / tiếp thu (1–5): ___ / ___",
            "Copy từ mới → Danh Sách Từ Vựng; lỗi tiêu biểu → Sổ Tay Lỗi Sai.",
            "1 dòng vào Sheet Theo Dõi Buổi Học (ghi rõ loại buổi A/B/C).",
        ],
    )
    out = BASE / "01 - Lộ Trình & Kế Hoạch" / "📝 [TEMPLATE] Giáo Án Từng Buổi.docx"
    doc.save(out)
    print("Wrote", out)


def build_bai_tap_tuan():
    doc = new_doc()
    doc.add_heading("[TEMPLATE] BÀI TẬP TUẦN ___ — NĂM 1 NỀN", 0)
    doc.add_paragraph(
        "Phù hợp Lộ Trình Năm 1: ưu tiên ngữ pháp + từ vựng + receptive skills; không bắt buộc format IELTS full."
    )
    add_block(
        doc,
        "THÔNG TIN",
        [
            "Tuần số:     ___    |  Tháng: ___",
            "Từ ngày:     ___/___/______  Đến: ___/___/______",
            "Chủ đề tuần / ngữ pháp trọng tâm: ___________________________",
            "Deadline nộp: trước buổi ___",
        ],
    )
    add_block(
        doc,
        "PHẦN A — TỪ VỰNG (mỗi ngày ~15 phút)",
        [
            "Học thuộc ___ từ từ buổi học + danh sách; hoàn thành: điền từ / đặt câu / collocation.",
        ],
    )
    add_block(
        doc,
        "PHẦN B — NGỮ PHÁP (mỗi ngày ~20 phút)",
        [
            "Làm đúng số câu GV ghi trong Giáo Án: ___ câu. Chép lại 3 câu SAI → SAI → ĐÚNG vào Sổ Tay.",
        ],
    )
    add_block(
        doc,
        "PHẦN C — KỸ NĂNG (chọn theo tuần; mỗi ngày ~20–25 phút)",
        [
            "Listening: 1 đoạn + bài tập kèm (link/ghi chú): ___________________________",
            "Reading: 1 đoạn + 5–8 câu hỏi (detail / main idea): ___________________________",
            "Writing: viết đoạn ___ từ chủ đề: ___________________________",
            "Speaking: ghi âm trả lời 3 câu hỏi (mỗi câu 45–90 giây), nộp link/ghi chú: ___________________________",
        ],
    )
    out = BASE / "03 - Bài Tập Theo Tuần" / "📝 [TEMPLATE] Bài Tập Từng Tuần.docx"
    doc.save(out)
    print("Wrote", out)


def build_test_mini():
    doc = new_doc()
    doc.add_heading("[TEMPLATE] KIỂM TRA NỀN MINI — TUẦN ___", 0)
    doc.add_paragraph(
        "Năm 1: kiểm tra nền ~30 phút mỗi tuần. Không dùng format IELTS — chỉ kiểm tra năng lực. "
        "Dữ liệu → Sheet Tổng Kết Tuần."
    )
    add_block(
        doc,
        "THÔNG TIN",
        ["Ngày: ___/___/______  Tuần: ___  Tháng: ___", "Thời gian: 30 phút (bấm giờ)."],
    )
    add_block(
        doc,
        "PHẦN 1 — NGỮ PHÁP (10 phút)",
        ["GV soạn 12–20 câu trắc nghiệm/tự luận ngắn theo chủ đề tuần.", "Điểm: ___/10"],
    )
    add_block(
        doc,
        "PHẦN 2 — LISTENING (8–10 phút)",
        ["1 đoạn 2–3 phút; 5–7 câu hỏi (không cần dạng IELTS).", "Điểm: ___/10"],
    )
    add_block(
        doc,
        "PHẦN 3 — READING (8–10 phút)",
        ["Đoạn 180–280 từ; câu hỏi detail/vocab in context.", "Điểm: ___/10"],
    )
    add_block(
        doc,
        "PHẦN 4 — VIẾT NGẮN (8–10 phút)",
        ["Viết 80–120 từ theo gợi ý; chấm theo rubric nền (grammar / task / coherence).", "Điểm: ___/10"],
    )
    doc.add_paragraph("TỔNG (gợi ý): ___/40 — quy đổi tỷ lệ khi nhập Sheet Tuần nếu cần.")
    out = BASE / "03 - Bài Tập Theo Tuần" / "📝 [TEMPLATE] Bài Test Mini Hàng Tuần.docx"
    doc.save(out)
    print("Wrote", out)


def build_kiem_tra_thang():
    doc = new_doc()
    doc.add_heading("[TEMPLATE] KIỂM TRA THÁNG ___ — NĂM 1 NỀN", 0)
    doc.add_paragraph(
        "Bài kiểm tra cuối tháng: đủ 4 kỹ năng + đánh giá phát âm; rubric là NĂNG LỰC NỀN. "
        "Năm 1 KHÔNG dùng format IELTS — chỉ kiểm tra năng lực tiếng Anh thuần. "
        "Định hướng Năm 2 (IELTS hay khác) chốt sau bài kiểm tra tháng 12."
    )
    add_block(
        doc,
        "THÔNG TIN",
        [
            "Tháng: ___  Ngày: ___/___/______",
            "Phạm vi: tuần ___ → ___",
            "Rubric chấm (ghi sẵn trên đề): Grammar / Vocab / Listening / Reading / Writing / Speaking (tỷ lệ GV tự chọn).",
        ],
    )
    add_block(
        doc,
        "PHẦN 1 — TỪ VỰNG (20 điểm gợi ý)",
        ["Chọn ~20 từ đã học trong tháng từ Sheet Từ Vựng; tránh ép thuần academic nếu tháng đầu nền."],
    )
    add_block(
        doc,
        "PHẦN 2 — NGỮ PHÁP (25 điểm gợi ý)",
        ["Kiểm tra các chủ đề đã dạy trong tháng (mixed exercises)."],
    )
    add_block(
        doc,
        "PHẦN 3 — ĐỌC (25 điểm gợi ý)",
        ["1–2 đoạn; câu hỏi nền (detail, paraphrase, main idea)."],
    )
    add_block(
        doc,
        "PHẦN 4 — VIẾT (20 điểm gợi ý)",
        ["120–200 từ theo chủ đề tháng; có outline hoặc không tùy tháng."],
    )
    add_block(
        doc,
        "PHẦN 5 — NGHE (10 điểm gợi ý)",
        ["1 đoạn + câu hỏi trắc nghiệm/ghi chú."],
    )
    add_block(
        doc,
        "PHẦN 6 — NÓI (10 điểm gợi ý)",
        ["2–3 phút có gợi ý bullet; chấm theo: trôi chảy, phát âm, ngữ pháp khi nói, ý."],
    )
    doc.add_paragraph(
        "Rubric 6 kỹ năng (thang 1-5) — không quy đổi band IELTS trong Năm 1. "
        "Tháng 12: thêm cột tự đánh giá của HV bên cạnh GV để chốt định hướng Năm 2."
    )
    out = BASE / "04 - Kiểm Tra & Đánh Giá" / "📝 [TEMPLATE] Bài Kiểm Tra Tháng.docx"
    doc.save(out)
    print("Wrote", out)


def build_bao_cao():
    doc = new_doc()
    doc.add_heading("[TEMPLATE] BÁO CÁO THÁNG GỬI HỌC VIÊN — NĂM 1 NỀN", 0)
    doc.add_paragraph(
        "Phụ thuộc: Sheet Đánh Giá Tháng (điểm/rubric nền); Tổng kết tuần; Bài kiểm tra tháng."
    )
    add_block(
        doc,
        "Email / Slack gửi HV",
        [
            "Gửi Huyền Trang,",
            "",
            "Dưới đây là báo cáo tháng ___ (Năm 1 — nền tảng).",
            "Số buổi đã học: ___ / 12 dự kiến.",
        ],
    )
    add_block(
        doc,
        "KẾT QUẢ KIỂM TRA THÁNG (rubric nền)",
        [
            "Từ vựng:    ___/___",
            "Ngữ pháp:   ___/___",
            "Đọc:        ___/___",
            "Viết:       ___/___",
            "Nghe:       ___/___",
            "Nói:        ___/___",
            "Nhận xét tổng thể (ưu tiên): ___________________________",
            "Chuyên đề phát âm tháng (kết quả): ___________________________",
        ],
    )
    add_block(
        doc,
        "NỘI DUNG ĐÃ HỌC",
        [
            "Ngữ pháp chính: ___________________________",
            "Từ vựng/chủ đề:  ___________________________",
            "Kỹ năng luyện nhiều nhất: ___________________________",
        ],
    )
    add_block(
        doc,
        "MỤC TIÊU THÁNG SAU (năng lực, không ép band)",
        ["___________________________________________________________"],
    )
    doc.add_paragraph("Mẹo: mở đầu bằng điểm tích cực; nêu 1–2 việc cần cải thiện cụ thể có hành động.")
    out = BASE / "04 - Kiểm Tra & Đánh Giá" / "📄 [TEMPLATE] Báo Cáo Tháng Gửi Học Viên.docx"
    doc.save(out)
    print("Wrote", out)


def main():
    build_giao_an()
    build_bai_tap_tuan()
    build_test_mini()
    build_kiem_tra_thang()
    build_bao_cao()


if __name__ == "__main__":
    main()
