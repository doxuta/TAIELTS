# -*- coding: utf-8 -*-
"""Rebuild [06] Tài nguyên + Nhật ký tiến bộ cho Năm 1 nền tiếng Anh thuần. Năm 2 để mở."""
from pathlib import Path

from docx import Document
from docx.shared import Pt

BASE = Path(__file__).resolve().parent.parent


def doc_header(title: str):
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)
    doc.add_heading(title, 0)
    return doc


def build_tai_nguyen():
    doc = doc_header("[06] TÀI NGUYÊN HỌC TẬP ĐƯỢC CHỌN LỌC — NĂM 1 NỀN TIẾNG ANH")
    doc.add_paragraph("Phụ thuộc: Hồ Sơ Học Viên. Năm 1: tài nguyên nền thuần — không IELTS. Năm 2: bổ sung tài nguyên theo định hướng đã chốt cuối Năm 1.")
    doc.add_paragraph(
        "Nguyên tắc: Ít mà chất — không lang thang. Cài Anki từ buổi đầu."
    )
    doc.add_heading("LISTENING", level=1)
    doc.add_paragraph("Tháng 1–2 (nếu HV đầu vào A1–A2):")
    doc.add_paragraph(
        "• VOA Learning English Level 1 (Beginner, ~1 phút) — https://learningenglish.voanews.com/p/5610.html\n"
        "  Nghe chậm, từ vựng cốt lõi; 3–4 đoạn/tuần. Có transcript đi kèm.",
        style="List Bullet",
    )
    doc.add_paragraph("Tháng 3–12 (B1 trở lên):")
    doc.add_paragraph(
        "• BBC 6 Minute English — https://www.bbc.co.uk/learningenglish/english/features/6-minute-english\n"
        "  Nghe 1 lần không script → nghe lại có script; 3–4 lần/tuần.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "• VOA Learning English Level 2 — https://learningenglish.voanews.com/z/4729 (Intermediate).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "• YouTube: BBC Learning English (kênh chính); EngVid (chọn video < 10 phút).",
        style="List Bullet",
    )
    doc.add_paragraph("Năm 2: bổ sung tài nguyên theo định hướng (Cambridge IELTS / TED Talks / podcast pháp lý…) — sẽ chốt cuối Năm 1.")

    doc.add_heading("READING", level=1)
    doc.add_paragraph("NĂM 1:")
    doc.add_paragraph(
        "• British Council Reading (Elementary/Pre-Int) — https://learnenglish.britishcouncil.org/skills/reading",
        style="List Bullet",
    )
    doc.add_paragraph("• Graded Readers Level 1–3 (sách theo cấp độ).", style="List Bullet")
    doc.add_paragraph(
        "NĂM 2: Cambridge IELTS passages; báo chí (Guardian/BBC) với bài ngắn có câu hỏi.",
        style="List Bullet",
    )

    doc.add_heading("WRITING", level=1)
    doc.add_paragraph(
        "NĂM 1: Viết câu–đoạn theo bài tập giáo viên; Grammarly (https://www.grammarly.com); Ludwig (https://ludwig.guru)."
    )
    doc.add_paragraph(
        "Năm 2: tài nguyên chuyên biệt theo định hướng (IELTS Liz/Simon nếu chọn IELTS; Legal Writing nếu chọn tiếng Anh pháp lý; v.v.) — chốt cuối Năm 1."
    )

    doc.add_heading("SPEAKING", level=1)
    doc.add_paragraph(
        "NĂM 1: Shadowing BBC 6 Minute; nói theo gợi ý bullet trong giờ học; ELSA Speak (phát âm)."
    )
    doc.add_paragraph("Năm 2: theo định hướng (cue card 2 phút nếu IELTS; thuyết trình mô phỏng nếu học thuật; trao đổi nghiệp vụ nếu pháp lý). Quyết định cuối Năm 1.")

    doc.add_heading("TỪ VỰNG", level=1)
    doc.add_paragraph("Anki — https://apps.ankiweb.net (bắt buộc).")
    doc.add_paragraph("Oxford Learner's — https://www.oxfordlearnersdictionaries.com")
    doc.add_paragraph(
        "AWL: dùng dần từ tháng 4–12 Năm 1 — https://www.victoria.ac.nz/lals/resources/academicwordlist (không nhồi hết một lúc)."
    )

    doc.add_heading("ỨNG DỤNG ĐIỆN THOẠI", level=1)
    doc.add_paragraph("Anki, ELSA Speak, Cambridge Dictionary, YouTube BBC Learning English.")

    doc.add_heading("SÁCH GỢI Ý", level=1)
    doc.add_paragraph("Ngữ pháp (Năm 1) — chọn theo trình độ đầu vào:")
    doc.add_paragraph(
        "• Essential Grammar in Use (Raymond Murphy, BÌA ĐỎ) — cho HV ở A1–A2 (mới bắt đầu / mất gốc).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "• English Grammar in Use (Raymond Murphy, BÌA XANH) — cho HV ở B1 trở lên (đã có nền).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Từ vựng: English Vocabulary in Use — Pre-Intermediate & Intermediate (Stuart Redman) cho cả năm.",
        style="List Bullet",
    )
    doc.add_paragraph("Năm 2: sách chuyên biệt theo định hướng — chốt cuối Năm 1.")

    doc.add_heading("LỊCH TỰ HỌC GỢI Ý (ngày không có lớp)", level=1)
    doc.add_paragraph("Sáng 15' Anki | Trưa 10' đọc ngắn | Tối 20' nghe + bài tập GV — tổng ~45 phút.")

    out = BASE / "06 - Tài Liệu Tham Khảo" / "📚 [06] Tài Nguyên Học Tập Được Chọn Lọc.docx"
    doc.save(out)
    print("Wrote", out)


def build_nhat_ky():
    doc = doc_header("NHẬT KÝ TIẾN BỘ — HUYEN TRANG (NĂM 1 NỀN)")
    doc.add_paragraph(
        "Phụ thuộc: Sheet Buổi / Sheet Tuần. Milestone theo QUÝ & năng lực — không bắt buộc band mỗi tháng."
    )
    doc.add_paragraph("Ngày bắt đầu: ___/___/______")
    doc.add_paragraph("Ghi chú đầu vào (mô tả năng lực nền — không bắt buộc quy đổi band): ___________________________")

    for q in range(1, 5):
        doc.add_heading(f"QUÝ {q} — MỐC NỀN (gợi ý)", level=1)
        doc.add_paragraph("Điền khi đạt (checkbox hoặc ngày):")
        if q == 1:
            doc.add_paragraph(
                "[ ] Xong thì cơ bản + câu hỏi + viết đoạn ngắn\n[ ] Nói 90s–2 phút chủ đề quen\n[ ] ~600 từ nền đầu tiên",
                style="List Bullet",
            )
        elif q == 2:
            doc.add_paragraph(
                "[ ] Passive / conditional 1 / relative clauses ổn\n[ ] Đọc 250–350 từ\n[ ] Viết 130–150 từ có topic sentence",
                style="List Bullet",
            )
        elif q == 3:
            doc.add_paragraph(
                "[ ] Liên kết đoạn; nói dài hơn có outline\n[ ] Đọc phân tích nhẹ",
                style="List Bullet",
            )
        else:
            doc.add_paragraph(
                "[ ] Kiểm tra tổng hợp cuối năm — đủ 4 kỹ năng + phát âm\n"
                "[ ] Rubric năng lực 6 kỹ năng (thang 1-5) — GV & HV tự đánh giá song song\n"
                "[ ] Họp chốt định hướng Năm 2: IELTS / pháp lý / học thuật / giao tiếp (xem file lộ trình)",
                style="List Bullet",
            )
        doc.add_paragraph("Ghi chú cảm xúc / lời HV / đột phá quý: _____________________________")

    doc.add_heading("NHẬT KÝ THEO TUẦN (copy cho mỗi tuần)", level=1)
    doc.add_paragraph(
        "Tuần ___ | Tháng ___\nBuổi 1/2/3: điểm nổi bật — khó khăn — HV nói gì — GV quan sát.\nTổng tuần: từ mới ___ | rubric nền (không cần band) ___\nCần cải thiện: ___________________________"
    )

    doc.add_heading("NĂM 2 — ghi sau khi chốt định hướng (cuối Q4 Năm 1)", level=1)
    doc.add_paragraph("Định hướng đã chọn: [ ] IELTS  [ ] Pháp lý  [ ] Học thuật  [ ] Giao tiếp  [ ] Khác: _______")
    doc.add_paragraph("Lý do: ___________________________")
    doc.add_paragraph("Mục tiêu cụ thể đầu Năm 2: ___________________________")
    doc.add_paragraph("Nếu chọn IELTS: mục tiêu band / ngày đăng ký thi: ___________________________")

    out = BASE / "02 - Theo Dõi Tiến Độ" / "📓 [02] Nhật Ký Tiến Bộ Học Viên.docx"
    doc.save(out)
    print("Wrote", out)


def main():
    build_tai_nguyen()
    build_nhat_ky()


if __name__ == "__main__":
    main()
